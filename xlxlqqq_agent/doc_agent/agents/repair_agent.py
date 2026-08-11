"""
DocGuard Agent - Repair Agent
===============================

职责：
1. 读取 Review Agent 产出的 ReviewIssue 列表
2. 对每个 auto_repairable=True 的 issue 执行修复动作（直接操作 _docx_reference）
3. 对 auto_repairable=False 的 issue 添加批注（高亮 + 红色说明文本），不修改内容
4. 同步更新 StructuredDocument（runs / paragraphs 的 text / style）以保持模型一致
5. 写入 state.repaired_document / repair_actions / repair_success

修复动作映射（IssueCategory → RepairType）：
  - CONTENT_TYPO / CONTENT_TERMINOLOGY / CONTENT_WRONG_WORD → REPLACE_TEXT
  - FORMAT_FONT                            → CHANGE_FONT
  - FORMAT_SIZE                            → CHANGE_SIZE
  - FORMAT_SPACING                         → CHANGE_LINE_SPACING
  - FORMAT_INDENT                          → CHANGE_INDENT
  - FORMAT_ALIGNMENT                       → CHANGE_ALIGNMENT
  - STRUCTURE_MISSING_SECTION              → ADD_COMMENT（仅标注，不自动添加章节）
  - FORMAT_HEADING_LEVEL                   → ADD_COMMENT（语义相关，不自动调整）
  - 其他                                    → ADD_COMMENT（兜底）

设计要点：
- 修复直接作用于 python-docx Document（_docx_reference），不重新构建 DOCX
- 每个修复动作独立 try/except，单个失败不影响其他修复
- 每个修复后调用 DocxAnnotator 在修改位置添加高亮 + 批注
- 同步更新 StructuredDocument 的 runs/paragraphs，保持数据模型一致
- ReviewIssue.suggested_fix / original_text 提供修复值（错别字替换/字体/字号等）
- RepairAction 记录原值/新值/执行结果/批注状态，便于追溯
- 不依赖 LLM，纯规则实现，支持无 API Key 环境

输入 state 字段：
- parsed_document: 必填（必须含 _docx_reference 才能修复）
- review_issues: 必填（驱动修复决策）
- style_profile: 可选（提供字体/字号规范值）

输出 state 字段：
- repaired_document: StructuredDocument（与 parsed_document 同一对象，已就地修改）
- repair_actions: list[RepairAction]
- repair_success: bool
- repair_error: Optional[str]
"""

from __future__ import annotations

import uuid
from dataclasses import dataclass, field
from typing import Any, Optional

from agents.base import BaseAgent
from core.config import DocGuardConfig
from core.exceptions import AgentError
from core.llm_client import LLMClient
from core.logging_config import get_logger
from core.state import (
    DocGuardState,
    Location,
    RepairAction,
    RepairConfirmation,
    ReviewIssue,
    StyleProfile,
)
from document.annotator import AnnotationConfig, DocxAnnotator
from document.models import (
    FontStyle,
    IssueCategory,
    Paragraph,
    ParagraphFormat,
    RepairType,
    Run,
    StructuredDocument,
)


logger = get_logger("agents.repair_agent")


# ============================================================
# 辅助函数
# ============================================================
def _action_id() -> str:
    return f"repair_{uuid.uuid4().hex[:12]}"


def _make_repair_action(
    issue: ReviewIssue,
    repair_type: RepairType,
    *,
    original_value: Optional[str] = None,
    new_value: Optional[str] = None,
    executed: bool = False,
    success: bool = False,
    error_message: Optional[str] = None,
    annotated: bool = False,
    skipped: bool = False,
    decision: Optional[str] = None,
) -> RepairAction:
    """构造 RepairAction。

    新增字段（Phase 6 HITL）：
    - skipped：是否被拒绝/跳过
    - decision："approve" / "reject" / "override"
    """
    return RepairAction(  # type: ignore[call-arg]
        action_id=_action_id(),
        issue_id=issue.get("issue_id", ""),
        repair_type=repair_type.value,
        location=issue.get("location", Location()),
        original_value=original_value,
        new_value=new_value,
        executed=executed,
        success=success,
        error_message=error_message,
        annotated=annotated,
        skipped=skipped,
        decision=decision or "approve",
    )


def _parse_suggested_fix(issue: ReviewIssue) -> Optional[str]:
    """
    从 issue.suggested_fix 中提取修复目标值。

    suggested_fix 形如：
      - "替换为『格式』"
      - "将字体修改为 宋体"
      - "将字号修改为 12.0pt"
      - "调整行距为 1.5"
      - "调整首行缩进为 24.0pt"
      - "替换为企业规范写法『PyTorch』"

    提取规则：取最后一个『』内的文本，无则返回 None。
    """
    fix = issue.get("suggested_fix")
    if not fix:
        return None
    # 优先匹配中文书名号
    if "『" in fix and "』" in fix:
        start = fix.rfind("『") + 1
        end = fix.rfind("』")
        if end > start:
            return fix[start:end]
    # 退化为取最后一个空格后的 token
    parts = fix.split()
    if len(parts) >= 2:
        return parts[-1]
    return None


# ============================================================
# 修复器基类
# ============================================================
@dataclass
class RepairResult:
    """单次修复结果"""
    success: bool = False
    original_value: Optional[str] = None
    new_value: Optional[str] = None
    error_message: Optional[str] = None
    annotated: bool = False


class BaseRepairer:
    """修复器基类，定义统一接口"""

    repair_type: RepairType = RepairType.ADD_COMMENT

    def __init__(self, annotator: Optional[DocxAnnotator] = None) -> None:
        self.annotator = annotator or DocxAnnotator()
        self.logger = get_logger(f"repair.{self.repair_type.value}")

    def repair(
        self,
        docx_doc: Any,
        structured_doc: StructuredDocument,
        issue: ReviewIssue,
    ) -> RepairResult:
        """
        执行修复。

        子类重写 _do_repair 实现具体逻辑；本方法统一处理异常与批注。

        Args:
            docx_doc: python-docx Document（_docx_reference）
            structured_doc: StructuredDocument（用于同步数据模型）
            issue: 待修复的 ReviewIssue

        Returns:
            RepairResult
        """
        result = RepairResult()
        try:
            result = self._do_repair(docx_doc, structured_doc, issue)
        except Exception as e:
            result.success = False
            result.error_message = str(e)
            self.logger.warning(
                "修复失败 | issue=%s | error=%s",
                issue.get("issue_id"), e, exc_info=True,
            )

        # 批注（无论修复成功与否，只要 executed=True 就标注）
        if result.success or result.annotated:
            try:
                self._annotate(docx_doc, issue, result)
                result.annotated = True
            except Exception as e:
                self.logger.warning("批注失败 | issue=%s | error=%s",
                                    issue.get("issue_id"), e)
        return result

    def _do_repair(
        self,
        docx_doc: Any,
        structured_doc: StructuredDocument,
        issue: ReviewIssue,
    ) -> RepairResult:
        raise NotImplementedError

    def _annotate(
        self,
        docx_doc: Any,
        issue: ReviewIssue,
        result: RepairResult,
    ) -> None:
        """在修改位置添加批注。

        容错：
        - PDF/PPT fallback 时 docx_doc 可能不是 python-docx Document（被
          RepairAgent 包装成 object() 占位），此时跳过批注避免 AttributeError。
        - para_idx 越界 / 段落不存在时捕获异常降级为无批注。
        """
        # 非真实 python-docx Document → 跳过（PDF/PPT 等 fallback 路径）
        if docx_doc is None or not hasattr(docx_doc, "paragraphs"):
            return

        loc = issue.get("location", {})
        para_idx = loc.get("paragraph_index")
        run_idx = loc.get("run_index")
        if para_idx is None:
            return

        try:
            paragraphs = docx_doc.paragraphs
            if not isinstance(paragraphs, (list, tuple)):
                return
            if para_idx >= len(paragraphs):
                # 段落索引越界（文档被修改过）→ 跳过批注不抛错
                return
        except Exception as e:
            self.logger.warning("段落访问异常，跳过批注: %s", e)
            return

        if result.success:
            comment = (
                f"已自动修复：{issue.get('title', '')}。"
                f"原文/原值: {result.original_value or '(空)'} → "
                f"新值: {result.new_value or '(空)'}"
            )
        else:
            comment = (
                f"需人工处理：{issue.get('title', '')}。"
                f"建议: {issue.get('suggested_fix', '(无)')}"
            )
        try:
            self.annotator.annotate_modification(
                docx_doc=docx_doc,
                paragraph_index=para_idx,
                run_index=run_idx,
                comment_text=comment,
                paragraph_id=loc.get("paragraph_id"),
            )
        except Exception as e:
            self.logger.warning(
                "批注写入失败（已跳过，不影响修复主流程）| issue=%s | error=%s",
                issue.get("issue_id"), e,
            )


# ============================================================
# 文本替换修复器（错别字 / 术语不一致）
# ============================================================
class TextReplaceRepairer(BaseRepairer):
    """替换段落文本中的错误写法为正确写法"""

    repair_type = RepairType.REPLACE_TEXT

    def _do_repair(
        self,
        docx_doc: Any,
        structured_doc: StructuredDocument,
        issue: ReviewIssue,
    ) -> RepairResult:
        loc = issue.get("location", {})
        para_idx = loc.get("paragraph_index")
        if para_idx is None:
            return RepairResult(error_message="缺少 paragraph_index")

        original = issue.get("original_text")
        new_value = _parse_suggested_fix(issue)
        if not original or not new_value:
            return RepairResult(error_message="缺少 original_text 或 suggested_fix")

        # 1. 操作 python-docx 段落
        if para_idx >= len(docx_doc.paragraphs):
            return RepairResult(error_message=f"段落索引越界: {para_idx}")

        docx_para = docx_doc.paragraphs[para_idx]
        replaced = False
        # 在 run 层面替换，保留格式
        for run in docx_para.runs:
            if original in run.text:
                run.text = run.text.replace(original, new_value)
                replaced = True
        # 若 run 层未命中（可能跨 run），则整段替换
        if not replaced:
            full_text = docx_para.text
            if original in full_text:
                # 整段替换会丢失 run 边界，作为兜底方案
                new_full = full_text.replace(original, new_value)
                # 清空所有 run，写入第一个 run
                if docx_para.runs:
                    first_run = docx_para.runs[0]
                    first_run.text = new_full
                    for r in docx_para.runs[1:]:
                        r.text = ""
                replaced = True

        if not replaced:
            return RepairResult(error_message=f"段落中未找到原文: {original}")

        # 2. 同步 StructuredDocument
        if para_idx < len(structured_doc.paragraphs):
            sd_para = structured_doc.paragraphs[para_idx]
            sd_para.text = sd_para.text.replace(original, new_value)
            for run in sd_para.runs:
                if original in run.text:
                    run.text = run.text.replace(original, new_value)

        return RepairResult(
            success=True,
            original_value=original,
            new_value=new_value,
        )


# ============================================================
# 字体修复器
# ============================================================
class FontChangeRepairer(BaseRepairer):
    """将段落/Run 的字体改为 style_profile 中的期望字体"""

    repair_type = RepairType.CHANGE_FONT

    def _do_repair(
        self,
        docx_doc: Any,
        structured_doc: StructuredDocument,
        issue: ReviewIssue,
    ) -> RepairResult:
        loc = issue.get("location", {})
        para_idx = loc.get("paragraph_index")
        run_idx = loc.get("run_index")
        if para_idx is None:
            return RepairResult(error_message="缺少 paragraph_index")

        # 从 suggested_fix 提取目标字体（如 "将字体修改为 宋体"）
        new_font = _parse_suggested_fix(issue)
        if not new_font:
            return RepairResult(error_message="无法从 suggested_fix 解析字体名")

        if para_idx >= len(docx_doc.paragraphs):
            return RepairResult(error_message=f"段落索引越界: {para_idx}")

        docx_para = docx_doc.paragraphs[para_idx]

        # 获取原值
        target_runs = []
        if run_idx is not None and run_idx < len(docx_para.runs):
            target_runs = [docx_para.runs[run_idx]]
        else:
            target_runs = list(docx_para.runs)

        if not target_runs:
            return RepairResult(error_message="段落无 Run 可修改")

        original_fonts = [r.font.name for r in target_runs]

        # 修改 python-docx Run
        from docx.oxml.ns import qn
        for run in target_runs:
            run.font.name = new_font
            # 同时设置东亚字体（中文显示需要）
            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.find(qn("w:rFonts"))
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                rPr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), new_font)
            rFonts.set(qn("w:ascii"), new_font)
            rFonts.set(qn("w:hAnsi"), new_font)

        # 同步 StructuredDocument
        if para_idx < len(structured_doc.paragraphs):
            sd_para = structured_doc.paragraphs[para_idx]
            for r in sd_para.runs:
                if run_idx is not None and r.run_index != run_idx:
                    continue
                r.style.name = new_font
                r.style.name_east_asian = new_font

        return RepairResult(
            success=True,
            original_value=original_fonts[0] if original_fonts else None,
            new_value=new_font,
        )


# ============================================================
# 字号修复器
# ============================================================
class SizeChangeRepairer(BaseRepairer):
    """将段落/Run 的字号改为期望值"""

    repair_type = RepairType.CHANGE_SIZE

    def _do_repair(
        self,
        docx_doc: Any,
        structured_doc: StructuredDocument,
        issue: ReviewIssue,
    ) -> RepairResult:
        loc = issue.get("location", {})
        para_idx = loc.get("paragraph_index")
        run_idx = loc.get("run_index")
        if para_idx is None:
            return RepairResult(error_message="缺少 paragraph_index")

        # 解析字号（如 "将字号修改为 12.0pt"）
        new_size = _parse_size(issue)
        if new_size is None:
            return RepairResult(error_message="无法从 suggested_fix 解析字号")

        if para_idx >= len(docx_doc.paragraphs):
            return RepairResult(error_message=f"段落索引越界: {para_idx}")

        docx_para = docx_doc.paragraphs[para_idx]
        target_runs = []
        if run_idx is not None and run_idx < len(docx_para.runs):
            target_runs = [docx_para.runs[run_idx]]
        else:
            target_runs = list(docx_para.runs)

        if not target_runs:
            return RepairResult(error_message="段落无 Run 可修改")

        from docx.shared import Pt
        original_sizes = [r.font.size.pt if r.font.size else None for r in target_runs]
        for run in target_runs:
            run.font.size = Pt(new_size)

        # 同步 StructuredDocument
        if para_idx < len(structured_doc.paragraphs):
            sd_para = structured_doc.paragraphs[para_idx]
            for r in sd_para.runs:
                if run_idx is not None and r.run_index != run_idx:
                    continue
                r.style.size_pt = new_size

        return RepairResult(
            success=True,
            original_value=str(original_sizes[0]) if original_sizes else None,
            new_value=f"{new_size}pt",
        )


# ============================================================
# 行距修复器
# ============================================================
class LineSpacingRepairer(BaseRepairer):
    """调整段落行距"""

    repair_type = RepairType.CHANGE_LINE_SPACING

    def _do_repair(
        self,
        docx_doc: Any,
        structured_doc: StructuredDocument,
        issue: ReviewIssue,
    ) -> RepairResult:
        loc = issue.get("location", {})
        para_idx = loc.get("paragraph_index")
        if para_idx is None:
            return RepairResult(error_message="缺少 paragraph_index")

        new_spacing = _parse_float(issue)
        if new_spacing is None:
            return RepairResult(error_message="无法从 suggested_fix 解析行距")

        if para_idx >= len(docx_doc.paragraphs):
            return RepairResult(error_message=f"段落索引越界: {para_idx}")

        docx_para = docx_doc.paragraphs[para_idx]
        original = (
            docx_para.paragraph_format.line_spacing
            if docx_para.paragraph_format.line_spacing is not None
            else None
        )

        # 设置行距：python-docx 会根据 float 值自动选择合适的 rule
        # （如 1.5 → ONE_POINT_FIVE，2.0 → DOUBLE，1.0 → SINGLE）
        # 若需精确控制，可显式设为 MULTIPLE 规则
        from docx.enum.text import WD_LINE_SPACING
        if new_spacing in (1.0, 1.5, 2.0):
            # 内置预设值，直接设置 line_spacing 让 python-docx 自动选 rule
            docx_para.paragraph_format.line_spacing = new_spacing
        else:
            # 任意倍数，使用 MULTIPLE 规则
            docx_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            docx_para.paragraph_format.line_spacing = new_spacing

        # 同步 StructuredDocument
        if para_idx < len(structured_doc.paragraphs):
            sd_para = structured_doc.paragraphs[para_idx]
            sd_para.style.line_spacing = new_spacing
            sd_para.style.line_spacing_rule = "multiple"

        return RepairResult(
            success=True,
            original_value=str(original) if original is not None else None,
            new_value=str(new_spacing),
        )


# ============================================================
# 首行缩进修复器
# ============================================================
class IndentChangeRepairer(BaseRepairer):
    """调整段落首行缩进"""

    repair_type = RepairType.CHANGE_INDENT

    def _do_repair(
        self,
        docx_doc: Any,
        structured_doc: StructuredDocument,
        issue: ReviewIssue,
    ) -> RepairResult:
        loc = issue.get("location", {})
        para_idx = loc.get("paragraph_index")
        if para_idx is None:
            return RepairResult(error_message="缺少 paragraph_index")

        new_indent = _parse_float(issue)
        if new_indent is None:
            return RepairResult(error_message="无法从 suggested_fix 解析缩进值")

        if para_idx >= len(docx_doc.paragraphs):
            return RepairResult(error_message=f"段落索引越界: {para_idx}")

        docx_para = docx_doc.paragraphs[para_idx]
        from docx.shared import Pt
        original = (
            docx_para.paragraph_format.first_line_indent.pt
            if docx_para.paragraph_format.first_line_indent is not None
            else None
        )
        docx_para.paragraph_format.first_line_indent = Pt(new_indent)

        # 同步 StructuredDocument
        if para_idx < len(structured_doc.paragraphs):
            sd_para = structured_doc.paragraphs[para_idx]
            sd_para.style.first_line_indent_pt = new_indent

        return RepairResult(
            success=True,
            original_value=str(original) if original is not None else None,
            new_value=f"{new_indent}pt",
        )


# ============================================================
# 批注修复器（仅标注，不修改内容）
# ============================================================
class CommentOnlyRepairer(BaseRepairer):
    """不修改内容，仅在问题位置添加批注"""

    repair_type = RepairType.ADD_COMMENT

    def _do_repair(
        self,
        docx_doc: Any,
        structured_doc: StructuredDocument,
        issue: ReviewIssue,
    ) -> RepairResult:
        # 标记为 annotated，但不修改内容
        return RepairResult(
            success=False,  # 未实际修复
            annotated=True,
            error_message="需人工处理（auto_repairable=False）",
        )


# ============================================================
# 解析辅助
# ============================================================
def _parse_size(issue: ReviewIssue) -> Optional[float]:
    """从 suggested_fix 中解析字号（如 '将字号修改为 12.0pt' → 12.0）"""
    fix = issue.get("suggested_fix")
    if not fix:
        return None
    import re
    m = re.search(r"(\d+(?:\.\d+)?)\s*pt", fix, re.IGNORECASE)
    if m:
        return float(m.group(1))
    return None


def _parse_float(issue: ReviewIssue) -> Optional[float]:
    """从 suggested_fix 中解析浮点数（如 '调整行距为 1.5' → 1.5）"""
    fix = issue.get("suggested_fix")
    if not fix:
        return None
    import re
    m = re.search(r"(\d+(?:\.\d+)?)", fix)
    if m:
        return float(m.group(1))
    return None


# ============================================================
# 修复器路由表
# ============================================================
def _get_repairer(
    issue: ReviewIssue,
    annotator: DocxAnnotator,
) -> BaseRepairer:
    """根据 issue.category 选择修复器"""
    cat = issue.get("category")
    auto = issue.get("auto_repairable", False)

    if not auto:
        return CommentOnlyRepairer(annotator)

    repairer_map = {
        IssueCategory.CONTENT_TYPO: TextReplaceRepairer,
        IssueCategory.CONTENT_TERMINOLOGY: TextReplaceRepairer,
        IssueCategory.CONTENT_WRONG_WORD: TextReplaceRepairer,
        IssueCategory.FORMAT_FONT: FontChangeRepairer,
        IssueCategory.FORMAT_SIZE: SizeChangeRepairer,
        IssueCategory.FORMAT_SPACING: LineSpacingRepairer,
        IssueCategory.FORMAT_INDENT: IndentChangeRepairer,
    }
    cls = repairer_map.get(cat, CommentOnlyRepairer)
    return cls(annotator)


# ============================================================
# Repair Agent
# ============================================================
class RepairAgent(BaseAgent):
    """
    Repair Agent：根据 ReviewIssue 执行自动修复 + 批注标记。

    依赖注入 DocxAnnotator 与可选的 style_profile（来自上游 Retrieval）。
    修复直接作用于 _docx_reference，无需 LLM。
    """

    def __init__(
        self,
        llm_client: Optional[LLMClient],
        config: DocGuardConfig,
        logger=None,
        *,
        annotator: Optional[DocxAnnotator] = None,
        annotator_config: Optional[AnnotationConfig] = None,
    ) -> None:
        super().__init__(
            llm_client, config,
            logger or get_logger("agents.repair_agent"),
        )
        self._annotator = annotator or DocxAnnotator(annotator_config)

    def agent_name(self) -> str:
        return "repair_agent"

    async def execute(self, state: DocGuardState) -> DocGuardState:
        """执行文档修复。

        Phase 6 HITL 集成：
        - 读取 state.repair_confirmations，若存在且某 issue 有对应确认：
          * approve → 正常修复
          * reject  → 跳过修复，仅写入 RepairAction（decision=reject）
          * override → 使用 override_text 作为替换，覆盖原 suggested_fix
        - 若无对应确认（即该 issue 无需 HITL）→ 按原逻辑修复。
        """
        parsed_doc = self._validate_state_field(state, "parsed_document")
        issues = self._validate_state_field(state, "review_issues")
        confirmations: list[RepairConfirmation] = state.get("repair_confirmations") or []
        confirm_map: dict[str, RepairConfirmation] = {
            c["issue_id"]: c for c in confirmations if c.get("issue_id")
        }

        # 校验 _docx_reference（修复必需，PDF/PPT 等格式 fallback 时为 None）
        docx_doc = parsed_doc._docx_reference
        source_fmt = getattr(parsed_doc, "source_format", "docx")
        if docx_doc is None and source_fmt == "docx":
            raise AgentError(
                "DOCX parsed_document 缺少 _docx_reference，无法执行修复"
                "（可能是从字典反序列化的文档）",
                context={"agent": self.agent_name()},
            )

        self.logger.info(
            "[Repair] 开始修复 | 文档=%s | 格式=%s | 问题数=%d | "
            "HITL 确认数=%d | 可自动修复=%d",
            parsed_doc.filename,
            source_fmt,
            len(issues),
            len(confirmations),
            sum(1 for i in issues if i.get("auto_repairable")),
        )

        actions: list[RepairAction] = []
        success_count = 0
        annotate_only_count = 0
        failure_count = 0
        skipped_count = 0

        for issue in issues:
            issue_id = issue.get("issue_id", "")
            conf = confirm_map.get(issue_id)
            decision = conf["decision"] if conf else "approve"

            if decision == "reject":
                # 用户拒绝修复：写一个 skipped action
                action = _make_repair_action(
                    issue,
                    RepairType.ADD_COMMENT,
                    executed=False,
                    success=False,
                    skipped=True,
                    error_message="HITL reject：用户拒绝修复",
                )
                action["decision"] = "reject"
                actions.append(action)
                skipped_count += 1
                continue

            # 处理 override：把 override_text 写入 issue 的临时副本
            active_issue = issue
            if decision == "override" and conf and conf.get("override_text"):
                active_issue = {
                    **issue,  # type: ignore[typeddict-item]
                    "suggested_fix": conf["override_text"],
                }

            repairer = _get_repairer(active_issue, self._annotator)
            try:
                if docx_doc is None:
                    # 非 DOCX 格式（PDF/PPT），无法直接回写，统一走 CommentOnly
                    # 这里用 base comment-only repairer 的 repair 方法
                    from document.annotator import DocxAnnotator
                    if not isinstance(repairer, CommentOnlyRepairer):
                        repairer = CommentOnlyRepairer(self._annotator)
                result = repairer.repair(
                    docx_doc if docx_doc is not None else object(),
                    parsed_doc,
                    active_issue,
                )
            except Exception as e:
                result = RepairResult(
                    success=False,
                    error_message=f"修复器异常: {e}",
                )
                self.logger.warning(
                    "[Repair] 修复器异常 | issue=%s | error=%s",
                    issue.get("issue_id"), e, exc_info=True,
                )

            action = _make_repair_action(
                issue,
                type(repairer).repair_type,
                original_value=result.original_value,
                new_value=result.new_value,
                executed=result.success,
                success=result.success,
                error_message=result.error_message,
                annotated=result.annotated,
            )
            if decision != "approve":
                action["decision"] = decision
            actions.append(action)

            if result.success:
                success_count += 1
            elif result.annotated:
                annotate_only_count += 1
            else:
                failure_count += 1

        # 写入 state（repaired_document 与 parsed_document 是同一对象，已就地修改）
        state["repaired_document"] = parsed_doc
        state["repair_actions"] = actions
        state["repair_success"] = (
            failure_count == 0 or success_count > 0 or len(actions) == 0
        )
        state["repair_error"] = None if failure_count == 0 else (
            f"{failure_count} 个修复动作失败"
        )

        self.logger.info(
            "[Repair] 修复完成 | 成功=%d | 仅批注=%d | 失败=%d | "
            "跳过(reject)=%d | 总动作=%d",
            success_count, annotate_only_count, failure_count,
            skipped_count, len(actions),
        )
        return state

    def _build_summary(self, state: DocGuardState) -> str:
        actions = state.get("repair_actions") or []
        success = sum(1 for a in actions if a.get("success"))
        annotated = sum(1 for a in actions if a.get("annotated"))
        return (
            f"actions={len(actions)}, success={success}, "
            f"annotated={annotated}, success_flag={state.get('repair_success')}"
        )
