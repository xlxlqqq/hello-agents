"""
DocGuard Agent - DOCX 解析器
==============================

设计要点：
1. 将 python-docx.Document 解析为 StructuredDocument（与 docx 库解耦的纯数据模型）
2. 完整提取：段落 / Run / 样式 / 标题层级 / 表格 / 图片 / 元数据
3. 单位转换：python-docx 使用 EMU 和 Twips，本解析器统一转为 Pt（磅）
4. 保留原始 docx.Document 引用（_docx_reference），供 Repair Agent 直接回写
5. 解析失败时抛出 DocumentParseError，携带上下文信息
6. 提供样式统计能力，便于规则匹配

python-docx 关键 API：
- doc.paragraphs: 段落列表（不含表格内段落）
- doc.tables: 表格列表
- doc.inline_shapes: 内联图片
- paragraph.style.name: 样式名（"Heading 1", "Normal" 等）
- paragraph.runs: Run 列表
- run.font: Font 对象（name/size/bold/italic/color）
"""

from __future__ import annotations

import logging
import os
from pathlib import Path
from typing import Any, Optional

from docx import Document
from docx.document import Document as _DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Emu, Pt, Twips

from core.exceptions import DocumentParseError, wrap_exception
from core.logging_config import get_logger
from document.models import (
    FontStyle,
    ImageInfo,
    Paragraph,
    ParagraphFormat,
    Run,
    StructuredDocument,
    Table,
    TableCell,
    TableRow,
)

logger = get_logger("document.parser")


# ============================================================
# 单位转换工具
# ============================================================
def emu_to_pt(emu: Optional[int]) -> Optional[float]:
    """EMU 转 Pt（1 Pt = 12700 EMU）"""
    if emu is None:
        return None
    return round(emu / 12700.0, 2)


def twips_to_pt(twips: Optional[int]) -> Optional[float]:
    """Twips 转 Pt（1 Pt = 20 Twips）"""
    if twips is None:
        return None
    return round(twips / 20.0, 2)


def alignment_to_str(alignment: Any) -> Optional[str]:
    """WD_ALIGN_PARAGRAPH 枚举转字符串"""
    if alignment is None:
        return None
    mapping = {
        WD_ALIGN_PARAGRAPH.LEFT: "left",
        WD_ALIGN_PARAGRAPH.CENTER: "center",
        WD_ALIGN_PARAGRAPH.RIGHT: "right",
        WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    }
    return mapping.get(alignment, str(alignment))


def line_spacing_rule_to_str(rule: Any) -> Optional[str]:
    """WD_LINE_SPACING 枚举转字符串"""
    if rule is None:
        return None
    mapping = {
        WD_LINE_SPACING.SINGLE: "single",
        WD_LINE_SPACING.ONE_POINT_FIVE: "1.5",
        WD_LINE_SPACING.DOUBLE: "double",
        WD_LINE_SPACING.AT_LEAST: "at_least",
        WD_LINE_SPACING.EXACTLY: "exact",
        WD_LINE_SPACING.MULTIPLE: "multiple",
    }
    return mapping.get(rule, str(rule))


def color_to_hex(color: Any) -> Optional[str]:
    """docx RGBColor 转十六进制字符串（无 # 前缀）"""
    if color is None:
        return None
    try:
        return str(color)
    except Exception:
        return None


# ============================================================
# 样式提取
# ============================================================
def extract_font_style(run: Any) -> FontStyle:
    """
    从 python-docx Run 对象提取字体样式。

    Args:
        run: docx.text.run.Run 实例

    Returns:
        FontStyle 实例
    """
    font = run.font
    style = FontStyle()

    # 字体名（西文）
    try:
        if font.name:
            style.name = font.name
    except Exception:
        pass

    # 东亚字体（中文）—— 需要从 rPr 的 w:rFonts 元素读取
    try:
        rpr = run._element.rPr
        if rpr is not None and rpr.rFonts is not None:
            east_asian = rpr.rFonts.get(qn("w:eastAsia"))
            if east_asian:
                style.name_east_asian = east_asian
    except Exception:
        pass

    # 字号
    try:
        if font.size is not None:
            style.size_pt = emu_to_pt(font.size)
    except Exception:
        pass

    # 加粗 / 斜体 / 下划线
    try:
        style.bold = font.bold if font.bold is not None else None
    except Exception:
        style.bold = None
    try:
        style.italic = font.italic if font.italic is not None else None
    except Exception:
        style.italic = None
    try:
        style.underline = font.underline if font.underline is not None else None
    except Exception:
        style.underline = None

    # 颜色
    try:
        if font.color and font.color.rgb is not None:
            style.color_hex = color_to_hex(font.color.rgb)
    except Exception:
        pass

    return style


def extract_paragraph_format(paragraph: Any) -> ParagraphFormat:
    """
    从 python-docx Paragraph 对象提取段落格式。

    Args:
        paragraph: docx.text.paragraph.Paragraph 实例

    Returns:
        ParagraphFormat 实例
    """
    pf = paragraph.paragraph_format
    fmt = ParagraphFormat()

    try:
        fmt.alignment = alignment_to_str(pf.alignment)
    except Exception:
        pass

    try:
        if pf.left_indent is not None:
            fmt.left_indent_pt = emu_to_pt(pf.left_indent)
    except Exception:
        pass

    try:
        if pf.right_indent is not None:
            fmt.right_indent_pt = emu_to_pt(pf.right_indent)
    except Exception:
        pass

    try:
        if pf.first_line_indent is not None:
            fmt.first_line_indent_pt = emu_to_pt(pf.first_line_indent)
    except Exception:
        pass

    try:
        if pf.line_spacing is not None:
            fmt.line_spacing = float(pf.line_spacing)
        fmt.line_spacing_rule = line_spacing_rule_to_str(pf.line_spacing_rule)
    except Exception:
        pass

    try:
        if pf.space_before is not None:
            fmt.space_before_pt = emu_to_pt(pf.space_before)
    except Exception:
        pass

    try:
        if pf.space_after is not None:
            fmt.space_after_pt = emu_to_pt(pf.space_after)
    except Exception:
        pass

    return fmt


# ============================================================
# 标题层级判定
# ============================================================
def detect_heading_level(paragraph: Any) -> Optional[int]:
    """
    判定段落的标题层级。

    优先使用样式名（"Heading 1" → 1），兜底使用 outline level。

    Args:
        paragraph: docx.text.paragraph.Paragraph 实例

    Returns:
        1-9 的标题层级，None 表示正文
    """
    # 方式1：样式名判定
    try:
        style_name = paragraph.style.name if paragraph.style else None
        if style_name:
            # "Heading 1" / "标题 1" / "Title"
            if style_name.lower().startswith("heading"):
                parts = style_name.split()
                if len(parts) >= 2 and parts[1].isdigit():
                    level = int(parts[1])
                    if 1 <= level <= 9:
                        return level
            # 中文 "标题 1" / "标题1"
            if "标题" in style_name:
                for part in style_name.split():
                    if part.isdigit():
                        level = int(part)
                        if 1 <= level <= 9:
                            return level
    except Exception:
        pass

    # 方式2：outline level（pPr 中的 w:outlineLvl）
    try:
        pPr = paragraph._element.pPr
        if pPr is not None and pPr.outlineLvl is not None:
            val = pPr.outlineLvl.get(qn("w:val"))
            if val is not None:
                level = int(val) + 1  # outlineLvl 是 0-based
                if 1 <= level <= 9:
                    return level
    except Exception:
        pass

    return None


def detect_list_info(paragraph: Any) -> tuple[bool, Optional[int], Optional[str]]:
    """
    检测段落是否是列表项。

    Args:
        paragraph: docx.text.paragraph.Paragraph 实例

    Returns:
        (is_list, list_level, list_style_name)
    """
    try:
        style_name = paragraph.style.name if paragraph.style else None
        if style_name and ("List" in style_name):
            # List Bullet / List Number / List Paragraph
            level = 0
            try:
                pPr = paragraph._element.pPr
                if pPr is not None and pPr.numPr is not None:
                    ilvl = pPr.numPr.ilvl
                    if ilvl is not None:
                        level = int(ilvl.get(qn("w:val"), "0"))
            except Exception:
                pass
            return True, level, style_name
    except Exception:
        pass

    # 检测 numPr（编号属性）
    try:
        pPr = paragraph._element.pPr
        if pPr is not None and pPr.numPr is not None:
            level = 0
            ilvl = pPr.numPr.ilvl
            if ilvl is not None:
                level = int(ilvl.get(qn("w:val"), "0"))
            return True, level, "List Number"
    except Exception:
        pass

    return False, None, None


# ============================================================
# 段落 / Run 解析
# ============================================================
def parse_run(run: Any, run_index: int) -> Run:
    """
    解析单个 python-docx Run 为 Run 模型。

    Args:
        run: docx.text.run.Run 实例
        run_index: 在段落中的索引

    Returns:
        Run 模型实例
    """
    return Run(
        text=run.text or "",
        style=extract_font_style(run),
        run_index=run_index,
    )


def parse_paragraph(
    paragraph: Any,
    paragraph_index: int,
    in_table: bool = False,
    parent_table_id: Optional[str] = None,
) -> Paragraph:
    """
    解析单个 python-docx Paragraph 为 Paragraph 模型。

    Args:
        paragraph: docx.text.paragraph.Paragraph 实例
        paragraph_index: 全局段落索引
        in_table: 是否在表格内
        parent_table_id: 所属表格 ID

    Returns:
        Paragraph 模型实例
    """
    # 解析所有 Run
    runs = [
        parse_run(run, run_idx)
        for run_idx, run in enumerate(paragraph.runs)
    ]

    # 标题层级
    heading_level = detect_heading_level(paragraph)

    # 列表信息
    is_list, list_level, list_style_name = detect_list_info(paragraph)

    # 样式名
    style_name = None
    try:
        if paragraph.style:
            style_name = paragraph.style.name
    except Exception:
        pass

    return Paragraph(
        text=paragraph.text or "",
        runs=runs,
        style=extract_paragraph_format(paragraph),
        style_name=style_name,
        heading_level=heading_level,
        is_list=is_list,
        list_level=list_level,
        list_style_name=list_style_name,
        in_table=in_table,
        parent_table_id=parent_table_id,
        paragraph_index=paragraph_index,
    )


# ============================================================
# 表格解析
# ============================================================
def parse_table(table: Any, table_index: int) -> Table:
    """
    解析 python-docx Table 为 Table 模型。

    Args:
        table: docx.table.Table 实例
        table_index: 表格在文档中的索引

    Returns:
        Table 模型实例
    """
    rows: list[TableRow] = []
    row_count = len(table.rows)
    col_count = len(table.columns) if row_count > 0 else 0

    for row_idx, row in enumerate(table.rows):
        cells: list[TableCell] = []
        for col_idx, cell in enumerate(row.cells):
            # 解析单元格中的段落
            cell_paragraphs = [
                parse_paragraph(
                    p,
                    paragraph_index=0,  # 单元格内段落索引单独维护
                    in_table=True,
                    parent_table_id=None,  # 稍后回填
                )
                for p in cell.paragraphs
            ]
            cells.append(TableCell(
                text=cell.text or "",
                paragraphs=cell_paragraphs,
                row_index=row_idx,
                col_index=col_idx,
                # row_span / col_span 较复杂，python-docx 不直接暴露，
                # 默认为 1，高级场景可从 vMerge 解析（暂不处理）
            ))
        rows.append(TableRow(
            row_index=row_idx,
            cells=cells,
            is_header=(row_idx == 0),  # 启发式：首行视为表头
        ))

    # 表格样式
    style_name = None
    try:
        if table.style:
            style_name = table.style.name
    except Exception:
        pass

    return Table(
        rows=rows,
        row_count=row_count,
        col_count=col_count,
        table_index=table_index,
        style_name=style_name,
        has_borders=True,  # 默认有边框，精确判断需查 tcBorders
    )


# ============================================================
# 图片解析
# ============================================================
def parse_images(doc: Any, paragraphs: list[Paragraph]) -> list[ImageInfo]:
    """
    提取文档中的内联图片信息。

    通过遍历 doc.inline_shapes 获取图片尺寸与所在段落。

    Args:
        doc: python-docx Document 实例
        paragraphs: 已解析的段落列表（用于关联段落）

    Returns:
        ImageInfo 列表
    """
    images: list[ImageInfo] = []
    try:
        for idx, shape in enumerate(doc.inline_shapes):
            # 查找包含该图片的段落
            paragraph_index = None
            paragraph_id = None
            try:
                # inline_shape 的 XML 元素
                inline_elem = shape._inline
                # 向上查找所在段落
                parent = inline_elem.getparent()
                while parent is not None and parent.tag != qn("w:p"):
                    parent = parent.getparent()
                if parent is not None:
                    # 在段落列表中匹配（通过元素引用）
                    for p_idx, para in enumerate(doc.paragraphs):
                        if para._element is parent:
                            paragraph_index = p_idx
                            if p_idx < len(paragraphs):
                                paragraph_id = paragraphs[p_idx].paragraph_id
                            break
            except Exception:
                pass

            # 内容类型
            content_type = None
            try:
                # 图片扩展名
                if hasattr(shape, "image"):
                    content_type = shape.image.content_type
            except Exception:
                pass

            images.append(ImageInfo(
                filename=f"image_{idx}",
                width_pt=emu_to_pt(shape.width),
                height_pt=emu_to_pt(shape.height),
                width_emu=shape.width,
                height_emu=shape.height,
                content_type=content_type,
                paragraph_id=paragraph_id,
                paragraph_index=paragraph_index,
            ))
    except Exception as e:
        logger.warning("图片解析失败（不影响主流程）: %s", e)

    return images


# ============================================================
# 元数据提取
# ============================================================
def extract_metadata(doc: Any) -> dict[str, Any]:
    """提取文档元数据（标题/作者/创建时间等）"""
    meta: dict[str, Any] = {
        "title": None,
        "author": None,
        "created": None,
        "modified": None,
    }
    try:
        cp = doc.core_properties
        meta["title"] = cp.title or None
        meta["author"] = cp.author or None
        meta["created"] = cp.created if cp.created else None
        meta["modified"] = cp.modified if cp.modified else None
    except Exception as e:
        logger.warning("元数据提取失败: %s", e)
    return meta


# ============================================================
# 样式统计
# ============================================================
def compute_style_stats(paragraphs: list[Paragraph]) -> tuple[dict[str, Any], dict[int, dict[str, Any]]]:
    """
    统计正文样式与各级标题样式，用于规则匹配。

    统计逻辑：取每个层级出现次数最多的样式作为"规范样式"。

    Args:
        paragraphs: 段落列表

    Returns:
        (body_style_stats, heading_style_stats)
        body_style_stats: {"font": ..., "size_pt": ..., "count": ...}
        heading_style_stats: {level: {"font": ..., "size_pt": ..., "count": ...}}
    """
    from collections import Counter

    body_fonts: Counter = Counter()
    body_sizes: Counter = Counter()
    body_count = 0

    heading_data: dict[int, dict[str, Counter]] = {}

    for p in paragraphs:
        if p.heading_level is not None:
            level = p.heading_level
            if level not in heading_data:
                heading_data[level] = {"fonts": Counter(), "sizes": Counter(), "count": 0}
            for run in p.runs:
                if run.style.name_east_asian:
                    heading_data[level]["fonts"][run.style.name_east_asian] += 1
                elif run.style.name:
                    heading_data[level]["fonts"][run.style.name] += 1
                if run.style.size_pt:
                    heading_data[level]["sizes"][run.style.size_pt] += 1
            heading_data[level]["count"] += 1
        elif p.text.strip():  # 非空正文段落
            body_count += 1
            for run in p.runs:
                if run.style.name_east_asian:
                    body_fonts[run.style.name_east_asian] += 1
                elif run.style.name:
                    body_fonts[run.style.name] += 1
                if run.style.size_pt:
                    body_sizes[run.style.size_pt] += 1

    # 取众数
    body_stats = {
        "font": body_fonts.most_common(1)[0][0] if body_fonts else None,
        "size_pt": body_sizes.most_common(1)[0][0] if body_sizes else None,
        "count": body_count,
    }

    heading_stats: dict[int, dict[str, Any]] = {}
    for level, data in heading_data.items():
        heading_stats[level] = {
            "font": data["fonts"].most_common(1)[0][0] if data["fonts"] else None,
            "size_pt": data["sizes"].most_common(1)[0][0] if data["sizes"] else None,
            "count": data["count"],
        }

    return body_stats, heading_stats


# ============================================================
# 主解析器
# ============================================================
class DocxParser:
    """
    DOCX 解析器：python-docx.Document → StructuredDocument。

    用法：
        parser = DocxParser()
        doc = parser.parse("input.docx")
        print(doc.get_statistics())
    """

    def __init__(self) -> None:
        self.logger = get_logger("document.parser")

    def parse(self, file_path: str) -> StructuredDocument:
        """
        解析 DOCX 文件为 StructuredDocument。

        Args:
            file_path: DOCX 文件路径

        Returns:
            StructuredDocument 实例（包含 _docx_reference）

        Raises:
            DocumentParseError: 文件不存在 / 格式错误 / 解析失败
        """
        path = Path(file_path)
        self.logger.info("开始解析 DOCX: %s", path)

        # 文件校验
        if not path.exists():
            raise DocumentParseError(
                f"文件不存在: {file_path}",
                context={"file_path": file_path},
            )
        if not path.is_file():
            raise DocumentParseError(
                f"路径不是文件: {file_path}",
                context={"file_path": file_path},
            )
        if path.suffix.lower() not in (".docx",):
            raise DocumentParseError(
                f"仅支持 .docx 格式，当前文件后缀: {path.suffix}",
                context={"file_path": file_path, "suffix": path.suffix},
            )

        # 打开文档
        try:
            docx_doc = Document(str(path))
        except Exception as e:
            raise wrap_exception(
                e,
                DocumentParseError,
                f"DOCX 文件打开失败（可能已损坏）: {e}",
                file_path=str(path),
            ) from e

        # 解析各部分
        try:
            paragraphs = self._parse_all_paragraphs(docx_doc)
            tables = self._parse_all_tables(docx_doc)
            images = parse_images(docx_doc, paragraphs)
            metadata = extract_metadata(docx_doc)
            body_stats, heading_stats = compute_style_stats(paragraphs)
        except Exception as e:
            raise wrap_exception(
                e,
                DocumentParseError,
                f"DOCX 内容解析失败: {e}",
                file_path=str(path),
            ) from e

        # 字数统计
        word_count = sum(len(p.text) for p in paragraphs)

        # 构建 StructuredDocument
        structured = StructuredDocument(
            file_path=str(path),
            filename=path.name,
            paragraphs=paragraphs,
            tables=tables,
            images=images,
            title=metadata["title"],
            author=metadata["author"],
            created=metadata["created"],
            modified=metadata["modified"],
            word_count=word_count,
            body_style_stats=body_stats,
            heading_style_stats=heading_stats,
            _docx_reference=docx_doc,
        )

        self.logger.info(
            "DOCX 解析完成: %s | 段落=%d | 表格=%d | 图片=%d | 字数=%d",
            path.name,
            len(paragraphs),
            len(tables),
            len(images),
            word_count,
        )

        return structured

    # ============================================================
    # 内部方法
    # ============================================================
    def _parse_all_paragraphs(self, doc: Any) -> list[Paragraph]:
        """解析所有顶层段落（不含表格内段落）"""
        paragraphs: list[Paragraph] = []
        for idx, para in enumerate(doc.paragraphs):
            paragraphs.append(parse_paragraph(para, paragraph_index=idx))
        return paragraphs

    def _parse_all_tables(self, doc: Any) -> list[Table]:
        """解析所有表格"""
        tables: list[Table] = []
        for idx, table in enumerate(doc.tables):
            tables.append(parse_table(table, table_index=idx))
        return tables


# ============================================================
# 便捷函数
# ============================================================
def parse_docx(file_path: str) -> StructuredDocument:
    """
    便捷函数：解析 DOCX 文件。

    Args:
        file_path: DOCX 文件路径

    Returns:
        StructuredDocument 实例

    Raises:
        DocumentParseError: 解析失败
    """
    return DocxParser().parse(file_path)
