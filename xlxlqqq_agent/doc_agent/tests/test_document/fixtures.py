"""
测试 fixture：构造用于测试的 DOCX 文件
"""

from __future__ import annotations

from pathlib import Path
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from core.logging_config import get_logger

logger = get_logger("tests.fixtures")


def set_run_font(run, font_name: str, size_pt: float, *, bold: bool = False) -> None:
    """设置 Run 的字体（含东亚字体）"""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    # 设置东亚字体
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


def create_sample_docx(
    output_path: Path,
    *,
    with_headings: bool = True,
    with_table: bool = True,
    with_format_issues: bool = True,  # 默认含错别字号和错别字，方便 Review/Repair 场景触发 issue
) -> Path:
    """
    创建一个用于测试的 DOCX 文件。

    Args:
        output_path: 输出路径
        with_headings: 是否包含标题层级
        with_table: 是否包含表格
        with_format_issues: 是否故意引入格式问题（如错误字号）和错别字。
                           True（默认）：字号=10pt，段落含错别字"格试"
                           False：字号=12pt（规范），段落"格式"正确写法，无错别字

    Returns:
        输出文件 Path
    """
    doc = Document()

    # ----- 标题 -----
    if with_headings:
        h1 = doc.add_heading("项目背景", level=1)
        h2 = doc.add_heading("需求概述", level=2)
        h3 = doc.add_heading("功能需求", level=3)

    # ----- 正文段落 -----
    p1 = doc.add_paragraph()
    run1 = p1.add_run("本文档描述了 DocGuard Agent 系统的设计方案。")
    if with_format_issues:
        # 故意使用错误字号（10pt，应为 12pt）
        set_run_font(run1, "宋体", 10.0)
    else:
        set_run_font(run1, "宋体", 12.0)
    p1.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p2 = doc.add_paragraph()
    run2 = p2.add_run("系统采用 LLM + RAG + Multi-Agent 架构。")
    set_run_font(run2, "宋体", 12.0)

    # 含错别字的段落（供 Review/Repair 测试）；默认含错别字（保证 Repair 单测 &
    # 各集成测试可检测 issue），只有显式 with_format_issues=False 时才干净。
    p3 = doc.add_paragraph()
    if with_format_issues is False:
        # 干净文档（供 ValidationAgent no-issue-pass 等场景）
        run3 = p3.add_run("这个系统能够自动检查文档中的错别字和格式问题。")
    else:
        # 故意加入错别字"格试"（应为"格式"），触发内容审查
        run3 = p3.add_run("这个系统能够自动检查文档中的错别字和格试问题。")
    set_run_font(run3, "宋体", 12.0)

    # ----- 表格 -----
    if with_table:
        table = doc.add_table(rows=2, cols=3)
        table.style = "Table Grid"
        # 表头
        headers = ["模块", "技术栈", "说明"]
        for i, h in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = h
        # 数据行
        data = [
            ["后端", "FastAPI", "RESTful API"],
            ["Agent", "LangGraph", "工作流编排"],
        ]
        # 仅写入一行数据（表格 2 行）
        for i, row_data in enumerate(data[:1]):
            for j, val in enumerate(row_data):
                table.rows[1].cells[j].text = val

    # ----- 列表 -----
    doc.add_paragraph("主要功能：", style="List Bullet")
    doc.add_paragraph("文档解析", style="List Bullet")
    doc.add_paragraph("格式检查", style="List Bullet")

    # 保存
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info("测试 DOCX 已生成: %s", output_path)
    return output_path


def create_empty_docx(output_path: Path) -> Path:
    """创建一个空 DOCX 文件（仅含一个空段落）"""
    doc = Document()
    doc.add_paragraph("")
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path


def create_complex_docx(output_path: Path) -> Path:
    """创建一个较复杂的 DOCX（多层标题 + 多表格 + 列表）"""
    doc = Document()

    # 多层标题
    doc.add_heading("系统设计文档", level=1)
    doc.add_heading("1 总体设计", level=2)
    doc.add_heading("1.1 架构概述", level=3)
    doc.add_paragraph("系统采用分层架构，包含展示层、业务层、数据层。")

    doc.add_heading("1.2 模块划分", level=3)
    doc.add_paragraph("系统分为以下模块：")
    doc.add_paragraph("用户管理", style="List Bullet")
    doc.add_paragraph("文档处理", style="List Bullet")
    doc.add_paragraph("知识库", style="List Bullet")

    doc.add_heading("2 详细设计", level=2)
    doc.add_heading("2.1 数据库设计", level=3)

    # 复杂表格
    table = doc.add_table(rows=4, cols=4)
    table.style = "Table Grid"
    headers = ["表名", "字段数", "索引", "说明"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    rows_data = [
        ["users", "8", "id, email", "用户表"],
        ["documents", "12", "id, user_id", "文档表"],
        ["reviews", "15", "id, doc_id", "审查记录表"],
    ]
    for r_idx, row_data in enumerate(rows_data, start=1):
        for c_idx, val in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = val

    doc.add_heading("2.2 接口设计", level=3)
    doc.add_paragraph("提供 RESTful API，支持文档上传、查询、审查等操作。")

    doc.add_heading("3 测试方案", level=2)
    doc.add_paragraph("采用单元测试 + 集成测试 + 端到端测试的策略。")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
